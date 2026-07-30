"""
Word (.docx) rebalance-memo builder for the Aapryl clone tool.

`build_portfolio_docx(ctx)` returns the bytes of an editable Word document that
mirrors the Xponance rebalance memo: a Manager Allocation table, a New Managers
section (firm write-ups), a blank Detailed Portfolio Change section, the market
cycle chart image, the Proposed Portfolio Exposures tables (characteristics,
region, sector), and the FactSet active-style risk table.

All tabular content is real text (editable in Word); only the market-cycle
chart is embedded as a PNG. The caller (app.py /export_portfolio_docx) does the
data assembly and passes a plain `ctx` dict — see that endpoint for the shape.
"""
import os
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PATH     = os.path.join(os.path.dirname(__file__), 'static', 'assets', 'xponance_logo.png')
# Full-width branded header banner (Xponance logo + 30 Years + swoosh). Drop a
# replacement PNG at this exact path to override the built-in one.
HEADER_BANNER = os.path.join(os.path.dirname(__file__), 'static', 'assets', 'xponance_header.png')
# Full-width footer band (rule + cities + web/social). Override at this path.
FOOTER_BAND   = os.path.join(os.path.dirname(__file__), 'static', 'assets', 'xponance_footer_band.png')
PAGE_W_IN     = 8.5    # Letter width; images bleed to the page edges
LEFT_MARGIN_IN = 0.9
BRAND_BLUE  = '1F4E79'   # table header fill
BAND_BLUE   = '2E75B6'   # section brand band
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK        = RGBColor(0x1E, 0x2D, 0x3D)


# ── Low-level OOXML helpers ─────────────────────────────────────────────────
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _table_borders(table, color='BFBFBF', sz='4'):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


def _set_cell(cell, text, *, bold=False, color=None, align='left', size=9,
              italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run('' if text is None else str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _para(doc, text='', *, bold=False, italic=False, underline=False, size=10,
          space_after=6, align='left', keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = keep_with_next
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.size = Pt(size)
        run.font.color.rgb = DARK
    return p


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True   # don't strand a heading at page bottom
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)   # brand navy
    # brand-blue rule under the heading instead of an underline
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '3'); bottom.set(qn('w:color'), BAND_BLUE)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def _header_row(table, labels):
    cells = table.rows[0].cells
    for c, lab in zip(cells, labels):
        _set_cell(c, lab, bold=True, color=WHITE,
                  align='center' if lab else 'left', size=9)
        _shade(c, BRAND_BLUE)


def _add_table(doc, n_cols, header_labels, col_align=None, widths=None):
    t = doc.add_table(rows=1, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_borders(t)
    _header_row(t, header_labels)
    return t


# Usable text width with 0.9" side margins on Letter (8.5 - 1.8).
USABLE_IN = 6.7


def _apply_widths(table, widths):
    """Pin column widths with a fixed layout so Word honours them instead of
    auto-fitting (the default, which crushes label columns)."""
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    # Fixed layout keys off the tblGrid column definitions (NOT per-cell width),
    # so set those directly (twips = inches * 1440).
    grid_cols = table._tbl.tblGrid.findall(qn('w:gridCol'))
    for i, gc in enumerate(grid_cols):
        if i < len(widths):
            gc.set(qn('w:w'), str(int(round(widths[i].inches * 1440))))
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = w
    # Repeat the header row on every page and never split a row across pages,
    # so multi-page tables stay readable.
    hdr_trPr = table.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); hdr_trPr.append(th)
    for row in table.rows:
        cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true')
        row._tr.get_or_add_trPr().append(cs)


NEG_RED    = RGBColor(0xB0, 0x00, 0x00)
MUTED      = RGBColor(0x9A, 0xA4, 0xB0)
ZEBRA_FILL = 'F5F7FA'


def _heat_fill(v, maxabs):
    """Diverging red→white→green hex fill for a value. Negative = red,
    positive = green, magnitude sets intensity (normalised by maxabs)."""
    if v is None or maxabs <= 0:
        return None
    t = max(-1.0, min(1.0, v / maxabs))
    base = (255, 255, 255)
    tgt = (0x63, 0xBE, 0x7B) if t >= 0 else (0xF8, 0x69, 0x6B)  # green / red
    a = abs(t)
    r = int(round(base[0] + (tgt[0] - base[0]) * a))
    g = int(round(base[1] + (tgt[1] - base[1]) * a))
    b = int(round(base[2] + (tgt[2] - base[2]) * a))
    return f'{r:02X}{g:02X}{b:02X}'


def _has_fill(cell):
    tcPr = cell._tc.tcPr
    return tcPr is not None and tcPr.find(qn('w:shd')) is not None


def _num_cell(cell, s, size=9):
    """Right-aligned signed number: negatives as red parentheses (accounting
    style), zeros/blanks muted. `s` is a preformatted string like '-5.64'."""
    s = (s or '').strip()
    if s in ('', '--', '—', '0.00', '+0.0%', '0.0%'):
        _set_cell(cell, '—', align='right', size=size, color=MUTED)
        return
    neg = s.startswith('-')
    if s and s[0] in '+-':
        s = s[1:]
    if neg:
        s = '(' + s + ')'
    _set_cell(cell, s, align='right', size=size, color=(NEG_RED if neg else None))


def _zebra(table, first_data_row=1):
    for i, row in enumerate(table.rows):
        if i < first_data_row or (i - first_data_row) % 2 == 0:
            continue
        for c in row.cells:
            if not _has_fill(c):
                _shade(c, ZEBRA_FILL)


def _short_bench(name):
    """Compact display label for a long benchmark/index name."""
    if not name:
        return 'Benchmark'
    n = str(name).strip()
    low = n.lower()
    sc = ' SC' if ('small' in low or ' sc' in low) else ''
    if 'ex-united states' in low or 'ex us' in low or 'ex-us' in low or 'ex usa' in low:
        return 'ACWI ex-US' + sc
    if 'eafe' in low:
        return 'EAFE' + ('+Canada' if 'canada' in low else '') + sc
    if 'all country world' in low or 'acwi' in low:
        return 'ACWI' + sc
    if 'world' in low:
        return 'MSCI World' + sc
    if 'emerging' in low or low.startswith('msci em'):
        return 'EM' + sc
    return n if len(n) <= 20 else n[:18] + '…'


# ── Section builders ────────────────────────────────────────────────────────
def _manager_allocation(doc, ctx):
    _heading(doc, 'Manager Allocation')
    t = _add_table(doc, 5, ['Manager', 'Current Weight', 'Change in Weight',
                            'Target Weight', 'Action'])
    # Client total row
    row = t.add_row().cells
    _set_cell(row[0], ctx['client_name'] + ' — Total', bold=True, size=9)
    _set_cell(row[1], ctx['totals']['current'], bold=True, align='right', size=9)
    _set_cell(row[2], ctx['totals']['change'],  bold=True, align='right', size=9)
    _set_cell(row[3], ctx['totals']['target'],  bold=True, align='right', size=9)
    _set_cell(row[4], '', bold=True, size=9)
    for c in row:
        _shade(c, 'DCE6F1')
    for m in ctx['manager_rows']:
        row = t.add_row().cells
        _set_cell(row[0], m['name'], size=9)
        _set_cell(row[1], m['current'], align='right', size=9)
        _num_cell(row[2], m['change'])
        _set_cell(row[3], m['target'],  align='right', size=9)
        _set_cell(row[4], m.get('action', ''), align='center', size=9)
    _apply_widths(t, [Inches(2.35), Inches(1.15), Inches(1.2), Inches(1.1), Inches(0.9)])
    _zebra(t, first_data_row=2)


def _new_managers(doc, ctx):
    _heading(doc, 'New Managers')
    new = ctx.get('new_managers') or []
    if not new:
        _para(doc, 'No new managers in this rebalance.', italic=True, size=10)
        return
    for nm in new:
        _para(doc, nm['name'], bold=True, underline=True, size=10, space_after=2)
        _para(doc, nm.get('description') or '', size=10, space_after=8)


def _detailed_change(doc, ctx):
    _heading(doc, 'Detailed Portfolio Change')
    _para(doc, '', size=10)  # intentionally blank for manual context


def _market_cycle(doc, ctx):
    png = ctx.get('market_cycle_png')
    if not png:
        return
    _heading(doc, 'Manager Positioning — Market Cycle')
    try:
        doc.add_picture(BytesIO(png), width=Inches(6.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        # A malformed/empty capture shouldn't sink the whole memo.
        _para(doc, '[Market cycle chart unavailable]', italic=True, size=10)


def _characteristics(doc, ctx):
    exp = ctx.get('exposures')
    _heading(doc, 'Proposed Portfolio Exposures')
    if not exp:
        _para(doc, 'Upload a FactSet Exposures file to populate this section.',
              italic=True, size=10)
        return
    bench_lbl = _short_bench(exp.get('benchmark_matched'))
    _para(doc, 'Portfolio Characteristics', bold=True, size=10, space_after=2, keep_with_next=True)
    t = _add_table(doc, 4, ['', 'Pre-Trade', 'Post-Trade', bench_lbl])
    for c in exp['characteristics']:
        row = t.add_row().cells
        _set_cell(row[0], c['label'], size=9)
        _set_cell(row[1], c['pre'], align='right', size=9)
        _set_cell(row[2], c['post'], align='right', size=9)
        _set_cell(row[3], c['bench'], align='right', size=9)
    _apply_widths(t, [Inches(2.86), Inches(1.28), Inches(1.28), Inches(1.28)])
    _zebra(t, first_data_row=1)


def _weights_table(doc, title, rows, bench_lbl):
    _para(doc, title, bold=True, size=10, space_after=2, keep_with_next=True)
    t = doc.add_table(rows=2, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _table_borders(t)
    # Two-row merged header: Pre-Trade / Post-Trade each span Wt + Diff.
    def hc(cell, text, size=9):
        _set_cell(cell, text, bold=True, color=WHITE, align='center', size=size)
        _shade(cell, BRAND_BLUE)
    hc(t.cell(0, 0).merge(t.cell(1, 0)), '')
    hc(t.cell(0, 1).merge(t.cell(0, 2)), 'Pre-Trade')
    hc(t.cell(0, 3).merge(t.cell(0, 4)), 'Post-Trade')
    hc(t.cell(0, 5).merge(t.cell(1, 5)), bench_lbl)
    for ci, lab in ((1, 'Port Wt'), (2, 'Diff'), (3, 'Port Wt'), (4, 'Diff')):
        hc(t.cell(1, ci), lab, size=8)
    for r in rows:
        cells = t.add_row().cells
        label = ('    ' + r['label']) if r.get('indent') else r['label']
        _set_cell(cells[0], label, bold=r.get('header', False), size=9)
        _set_cell(cells[1], r['pre'],      align='right', size=9, bold=r.get('header', False))
        _num_cell(cells[2], r['pre_diff'])
        _set_cell(cells[3], r['post'],     align='right', size=9, bold=r.get('header', False))
        _num_cell(cells[4], r['post_diff'])
        _set_cell(cells[5], r['bench'],    align='right', size=9, bold=r.get('header', False))
        if r.get('header'):
            for c in cells:
                _shade(c, 'EAF1F8')
    _apply_widths(t, [Inches(1.7), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)])
    # Repeat BOTH header rows across page breaks.
    for hr in (t.rows[0], t.rows[1]):
        trPr = hr._tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None:
            th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)
    _zebra(t, first_data_row=2)


def _exposure_weight_tables(doc, ctx):
    exp = ctx.get('exposures')
    if not exp:
        return
    bench_lbl = _short_bench(exp.get('benchmark_matched'))
    _weights_table(doc, 'Region / Country Weights', exp['regions'], bench_lbl)
    _weights_table(doc, 'Sector Weights', exp['sectors'], bench_lbl)


def _risk_table(doc, ctx):
    risk = ctx.get('risk')
    _heading(doc, 'FactSet Active Style Exposures')
    if not risk or risk.get('error') or not risk.get('factors'):
        _para(doc, 'Upload a FactSet Risk Summary file to populate this section.',
              italic=True, size=10)
        return
    bench = (risk.get('benchmark') or {}).get('matched_column')
    if bench:
        _para(doc, f'Active exposures vs {bench}', italic=True, size=9, space_after=2)
    t = _add_table(doc, 3, ['Factor', 'Pre-Trade', 'Post-Trade'])
    cur, prop = risk.get('current', {}), risk.get('proposed', {})
    # Diverging heat scale: normalise to the largest absolute exposure across
    # both columns so the strongest reading saturates.
    vals = [v for v in list(cur.values()) + list(prop.values()) if v is not None]
    maxabs = max((abs(v) for v in vals), default=0.0) or 1.0
    for f in risk['factors']:
        row = t.add_row().cells
        _set_cell(row[0], f, size=9)
        for ci, val in ((1, cur.get(f)), (2, prop.get(f))):
            _set_cell(row[ci], f'{val:+.2f}' if val is not None else '—',
                      align='right', size=9)
            fill = _heat_fill(val, maxabs)
            if fill:
                _shade(row[ci], fill)
    _apply_widths(t, [Inches(2.7), Inches(2.0), Inches(2.0)])


def _signoff(doc, ctx):
    _para(doc, 'The precise rebalancing amounts for each manager will be finalized '
               'immediately prior to the trading day.', size=10, space_after=10)
    _para(doc, 'Kind regards,', size=10, space_after=24)
    _para(doc, '_________________________', size=10, space_after=0)
    _para(doc, 'Xponance Investment Team', size=10, space_after=0)


FOOT_GREY  = RGBColor(0x7F, 0x7F, 0x7F)
FOOT_BLUE  = RGBColor(0x4F, 0x9B, 0xD5)


def _page_number_field(paragraph):
    """Append a live PAGE-number field to a paragraph."""
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18'); rPr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '7F7F7F'); rPr.append(col)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = '1'; r.append(t); fld.append(r)
    paragraph._p.append(fld)


def _build_footer(section):
    """Footer: right-aligned page number, then the full-bleed cities/web band."""
    section.footer.is_linked_to_previous = False
    # Line 1: page number, right-aligned at the text margin (above the rule).
    p_num = section.footer.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.paragraph_format.space_before = Pt(0)
    p_num.paragraph_format.space_after = Pt(2)
    _page_number_field(p_num)
    # Line 2: full-bleed footer band image (rule + cities + web/social).
    if os.path.exists(FOOTER_BAND):
        bp = section.footer.add_paragraph()
        bp.paragraph_format.left_indent = Inches(-LEFT_MARGIN_IN)   # bleed to page edge
        bp.paragraph_format.space_before = Pt(0); bp.paragraph_format.space_after = Pt(0)
        try:
            bp.add_run().add_picture(FOOTER_BAND, width=Inches(PAGE_W_IN))
        except Exception:
            pass
    else:
        p = section.footer.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(USABLE_IN), WD_TAB_ALIGNMENT.RIGHT)
        left = p.add_run('Philadelphia, PA   |   Durham, NC')
        left.font.size = Pt(9); left.font.color.rgb = FOOT_GREY
        right = p.add_run('\txponance.com   |   @xponance   |   Xponance')
        right.font.size = Pt(9); right.font.color.rgb = FOOT_BLUE


# ── Entry point ─────────────────────────────────────────────────────────────
def build_portfolio_docx(ctx):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)
        # Full-bleed branded banner in the running page header.
        s.header_distance = Inches(0.3)
        s.footer_distance = Inches(0.3)
        if os.path.exists(HEADER_BANNER):
            hp = s.header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hp.paragraph_format.left_indent = Inches(-LEFT_MARGIN_IN)   # bleed to page edge
            hp.paragraph_format.space_after = Pt(2)
            try:
                hp.add_run().add_picture(HEADER_BANNER, width=Inches(PAGE_W_IN))
            except Exception:
                pass
        elif os.path.exists(LOGO_PATH):
            hp = s.header.paragraphs[0]
            hp.add_run().add_picture(LOGO_PATH, width=Inches(1.5))
        # Cities / website band footer
        _build_footer(s)

    # Memo header block
    _para(doc, f"Date: {ctx.get('date', '')}", bold=True, space_after=4)
    _para(doc, "To:", bold=True, space_after=4)
    _para(doc, "Cc:", bold=True, space_after=4)
    _para(doc, "From: Xponance Investment Team", bold=True, space_after=4)
    _para(doc, "Re:   Portfolio Rebalancing", bold=True, space_after=8)
    _para(doc, "The summary of our total proposed portfolio changes is below:",
          size=10, space_after=6)

    _manager_allocation(doc, ctx)
    _new_managers(doc, ctx)
    _detailed_change(doc, ctx)
    _market_cycle(doc, ctx)
    _characteristics(doc, ctx)
    _exposure_weight_tables(doc, ctx)
    _risk_table(doc, ctx)
    _signoff(doc, ctx)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
